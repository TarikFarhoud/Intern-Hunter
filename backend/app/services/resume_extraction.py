from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path
from zipfile import ZipFile, BadZipFile


MAX_EXTRACTED_TEXT_LEN = 200_000


def _normalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = value.replace("\u00a0", " ")
    value = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    value = re.sub(r"[ \t]{2,}", " ", value)
    value = value.strip()
    if len(value) > MAX_EXTRACTED_TEXT_LEN:
        value = value[:MAX_EXTRACTED_TEXT_LEN]
    return value


def _clean_fallback_text(value: str) -> str:
    """Remove common binary/container noise from best-effort DOC extraction.

    Legacy .doc fallbacks can produce a mix of real resume text plus long "garbage tail"
    (fonts, embedded objects, author/template fields). This cleaner is intentionally
    conservative: it keeps human-readable lines, but aggressively drops compact
    symbol-heavy lines and truncates after a long run of junk.
    """

    blocked_patterns = (
        "[content_types].xml",
        "_rels/.rels",
        "word/_rels/",
        "theme/theme",
        "docprops/",
        "application/vnd.openxmlformats",
        # Common Word binary metadata / templates / identifiers
        "normal.dotm",
        "microsoft office word",
        "microsoft word",
        "word.document",
        "msworddoc",
        "word.document.8",
        "_pid_",
        "_pid_hlinks",
        # Frequently leaked embedded/font tool markers
        "ttfautohint",
        "wrd_embed",
        "compobj",
    )

    blocked_exact = {
        # Font names that frequently leak from legacy .doc binaries
        "arial",
        "calibri",
        "cambria",
        "cambria math",
        "courier new",
        "georgia",
        "malgun gothic",
        "noto sans symbols",
        "symbol",
        "times new roman",
        "wingdings",
        # Common short binary markers
        "bjbj",
    }

    email_re = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
    url_re = re.compile(r"(?:https?://|www\.)\S+", re.IGNORECASE)
    phone_re = re.compile(r"\+?\d[\d\s().-]{7,}\d")

    kept: list[str] = []
    metadata_score = 0
    tail_bad_streak = 0
    seen_good = False

    for line in value.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        lowered = stripped.lower()
        if any(pattern in lowered for pattern in blocked_patterns) or lowered in blocked_exact:
            metadata_score = min(metadata_score + 2, 10)
            tail_bad_streak += 1
            if seen_good and tail_bad_streak >= 120:
                break
            continue

        has_email = bool(email_re.search(stripped))
        has_url = bool(url_re.search(stripped))
        has_phone = bool(phone_re.search(stripped))
        is_contact = has_email or has_url or has_phone

        alpha_count = sum(ch.isalpha() for ch in stripped)
        printable_count = sum(ch.isprintable() for ch in stripped)
        symbol_count = sum(not ch.isalnum() and not ch.isspace() for ch in stripped)
        alnum_space_count = sum(ch.isalnum() or ch.isspace() for ch in stripped)
        alnum_space_ratio = alnum_space_count / max(len(stripped), 1)
        symbol_ratio = symbol_count / max(len(stripped), 1)

        has_spaces = " " in stripped
        is_heading_like = bool(re.fullmatch(r"[A-Z][A-Z\s/&-]{2,}", stripped))
        word_token_count = len(re.findall(r"[A-Za-z]{2,}", stripped))

        # Compact, symbol-heavy, no-space lines are almost always binary noise.
        compact_symbol_heavy = (
            not has_spaces
            and not is_heading_like
            and not is_contact
            and len(stripped) <= 50
            and (symbol_ratio >= 0.15 or alnum_space_ratio < 0.85)
        )

        # Token-poor short lines (often random bytes decoded) are junk.
        token_poor_short = (
            not has_spaces
            and not is_heading_like
            and not is_contact
            and len(stripped) <= 16
            and (word_token_count == 0 or alpha_count < 4)
            and not stripped.isalpha()
        )

        looks_like_binary_token = (
            not has_spaces
            and not is_heading_like
            and not is_contact
            and len(stripped) >= 8
            and bool(re.fullmatch(r"[A-Za-z0-9`'\[\]{}()<>!@#$%^&*_=+\\|:;.,?-]{8,}", stripped))
        )

        low_quality = compact_symbol_heavy or token_poor_short or looks_like_binary_token
        if low_quality:
            metadata_score = min(metadata_score + 1, 10)
            tail_bad_streak += 1
            if seen_good and tail_bad_streak >= 120:
                break
            continue

        # Reset tail streak once we see a readable line.
        tail_bad_streak = 0

        # If we're in/near a metadata block, drop author/template labels.
        name_like = bool(re.fullmatch(r"[A-Z][a-z]{1,40},\s+[A-Z][a-z]{1,40}", stripped))
        label_like = lowered in {"title", "author", "subject", "company", "manager"}
        if metadata_score >= 4 and (name_like or label_like):
            continue

        if alpha_count < 2 and not is_contact and not is_heading_like:
            metadata_score = min(metadata_score + 1, 10)
            continue

        if printable_count and (alpha_count / max(len(stripped), 1)) < 0.2 and not is_contact and not is_heading_like:
            metadata_score = min(metadata_score + 1, 10)
            continue

        if symbol_count > len(stripped) * 0.4 and not is_contact and not is_heading_like:
            metadata_score = min(metadata_score + 1, 10)
            continue

        kept.append(stripped)
        seen_good = True
        metadata_score = max(metadata_score - 1, 0)

    # Trim any trailing font-only lines that slipped through.
    while kept and kept[-1].strip().lower() in blocked_exact:
        kept.pop()

    return "\n".join(kept)


def _extract_docx_xml_from_zip(file_path: Path) -> str | None:
    """If a file is actually a DOCX-like ZIP container, extract clean text from XML."""

    try:
        with ZipFile(file_path, "r") as archive:
            candidates = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
            if not candidates:
                return None

            parts: list[str] = []
            for name in candidates:
                try:
                    xml_text = archive.read(name).decode("utf-8", errors="ignore")
                except Exception:
                    continue

                # extract Word text runs from <w:t> ... </w:t>
                chunks = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml_text, flags=re.IGNORECASE | re.DOTALL)
                for chunk in chunks:
                    cleaned = (
                        chunk.replace("&amp;", "&")
                        .replace("&lt;", "<")
                        .replace("&gt;", ">")
                        .replace("&quot;", '"')
                        .replace("&#39;", "'")
                    )
                    if cleaned:
                        parts.append(cleaned)

            raw = "\n".join(parts)
            text = _normalize_text(raw)
            return text or None
    except BadZipFile:
        return None
    except Exception:
        return None


def _extract_doc_text_fallback(file_path: Path) -> str | None:
    """Best-effort text recovery for legacy .doc files without external tools.

    This is intentionally heuristic and may return partial text, but provides
    a usable preview when antiword/catdoc are not available.
    """

    try:
        data = file_path.read_bytes()
    except Exception:
        return None

    # Some files are DOCX containers with wrong extension (.doc).
    if data.startswith(b"PK"):
        text = _extract_docx_xml_from_zip(file_path)
        if text:
            return text

    ascii_chunks = re.findall(rb"[\x20-\x7e]{4,}", data)
    utf16_chunks = re.findall(rb"(?:[\x20-\x7e]\x00){4,}", data)

    parts: list[str] = []

    for chunk in ascii_chunks:
        text = chunk.decode("latin-1", errors="ignore")
        if text:
            parts.append(text)

    for chunk in utf16_chunks:
        text = chunk.decode("utf-16le", errors="ignore")
        if text:
            parts.append(text)

    if not parts:
        return None

    raw = "\n".join(parts)
    text = _clean_fallback_text(_normalize_text(raw))
    return text or None


def extract_resume_text(file_path: Path) -> str | None:
    """Extract text from a resume file.

    Supported:
    - .pdf via pdfminer.six
    - .docx via python-docx
    """

    ext = file_path.suffix.lower()

    try:
        if ext == ".pdf":
            from pdfminer.high_level import extract_text  # type: ignore[import-not-found]

            raw = extract_text(str(file_path)) or ""
            text = _normalize_text(raw)
            return text or None

        if ext == ".docx":
            from docx import Document  # type: ignore[import-not-found]

            doc = Document(str(file_path))
            parts: list[str] = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)

            raw = "\n".join(parts)
            text = _normalize_text(raw)
            return text or None

        if ext == ".doc":
            # Legacy Word format (.doc). Prefer a small external tool rather than
            # adding heavy Python dependencies.
            antiword = shutil.which("antiword")
            if antiword:
                proc = subprocess.run(
                    [antiword, str(file_path)],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=20,
                    check=False,
                )
                if proc.returncode == 0:
                    raw = proc.stdout.decode("utf-8", errors="ignore")
                    text = _normalize_text(raw)
                    if text:
                        return text

            # Fallback for environments without antiword (common on Windows dev).
            return _extract_doc_text_fallback(file_path)

    except Exception:
        return None

    return None
